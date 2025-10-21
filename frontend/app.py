# TFML Agentic AI — Luxe Console (Full App)
# -----------------------------------------
# - Dashboard: KPIs, alerts, charts, activity
# - Tenders: Filters + List/Kanban/Calendar + "Generate Draft Response"
# - Drafts Workspace: list, rich editor, attachments, lifecycle
# - Buttons styled for visibility (light & dark)
# - No audio control in Drafts (uploader restricted to docs/images only)
# - 6 seeded tenders with first drafts

import os
import re
import json
import time
from datetime import datetime, timedelta, date
from pathlib import Path
import sqlite3
import streamlit as st
import pandas as pd
import altair as alt
from docx import Document
from PIL import Image

# ======================================
# PATHS / CONFIG
# ======================================
try:
    BASE_DIR = Path(__file__).resolve().parent
except NameError:
    BASE_DIR = Path.cwd()

ASSETS = BASE_DIR / "assets"
LOGS = BASE_DIR / "logs"
EOIS = BASE_DIR / "eois"
TENDERS_DB = LOGS / "tenders.db"
LOGO_PATH = ASSETS / "tfml_logo.png"

EOIS.mkdir(parents=True, exist_ok=True)
LOGS.mkdir(parents=True, exist_ok=True)

# Brand
ACCENT = "#E60F18"
CARD = "#FFFFFF"
TEXT = "#000000"
MUTED = "#555555"
APP_BG_LIGHT = "#F9F9F9"
APP_BG_DARK = "#1E1E1E"
CARD_DARK = "#2A2A2A"

st.set_page_config(
    page_title="TFML Agentic AI — Luxe Console",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ======================================
# CSS (layout, tabs, buttons, uploader)
# ======================================
st.markdown(f"""
<style>
/* Base */
.stApp {{ background:{APP_BG_LIGHT}; color:{TEXT}; }}
.block-container {{ padding-top: 1rem; }}

/* Header */
.header {{ display:flex; align-items:center; gap:14px; padding:6px 0 14px; border-bottom:1px solid #ddd; }}
.header .title {{ font-weight:900; font-size:26px; color:{ACCENT}; letter-spacing:.2px; }}

/* Cards & Pills */
.kpi {{ background:{CARD}; border:1px solid #ddd; border-radius:14px; padding:16px; color:{TEXT}; }}
.kpi .label {{ color:{MUTED}; font-size:.78rem; text-transform:uppercase; letter-spacing:1px; }}
.kpi .value {{ font-size:1.8rem; font-weight:800; color:{ACCENT}; }}
.card {{ background:{CARD}; border:1px solid #ddd; border-radius:14px; padding:16px; color:{TEXT}; }}
.pill {{ display:inline-block; padding:2px 10px; border-radius:999px; font-size:.75rem; font-weight:700; background:#eee; color:{ACCENT}; border:1px solid {ACCENT}; }}

/* Tabs */
.stTabs [role="tablist"] {{ gap:8px; border-bottom:0; margin-bottom:.5rem; }}
.stTabs [role="tab"] {{
  padding:10px 16px; border:1px solid #ddd; background:#fff; color:#000;
  border-top-left-radius:10px; border-top-right-radius:10px; font-weight:700; opacity:1;
}}
.stTabs [role="tab"][aria-selected="true"] {{ background:{ACCENT}; color:#fff; border-color:{ACCENT}; }}

/* Dark Mode */
.dark-mode .stApp {{ background:{APP_BG_DARK}; color:#fff; }}
.dark-mode .kpi, .dark-mode .card {{ background:{CARD_DARK}; border-color:#444; color:#fff; }}
.dark-mode .header .title {{ color:#fff; }}
.dark-mode .pill {{ background:#444; color:#fff; border-color:#fff; }}
.dark-mode .stTabs [role="tab"] {{ background:{CARD_DARK}; border-color:#444; color:#fff; }}
.dark-mode .stTabs [role="tab"][aria-selected="true"] {{ background:{ACCENT}; border-color:{ACCENT}; color:#fff; }}

/* Buttons: visible text in light/dark, hover & disabled */
:root {{
  --tfml-btn-bg: {ACCENT};
  --tfml-btn-bg-hover: #c40e15;
  --tfml-btn-text: #ffffff;
  --tfml-btn-border: #9f0a10;
}}
.stButton > button,
.stDownloadButton > button,
form [data-testid="baseButton-secondary"],
form [data-testid="baseButton-primary"] {{
  background: var(--tfml-btn-bg) !important;
  color: var(--tfml-btn-text) !important;
  border: 1px solid var(--tfml-btn-border) !important;
  border-radius: 12px !important;
  font-weight: 700 !important;
  box-shadow: none !important;
}}
.stButton > button:hover,
.stDownloadButton > button:hover,
form [data-testid="baseButton-secondary"]:hover,
form [data-testid="baseButton-primary"]:hover {{
  background: var(--tfml-btn-bg-hover) !important;
  border-color: var(--tfml-btn-bg-hover) !important;
}}
.stButton > button:disabled,
.stDownloadButton > button:disabled,
form [data-testid="baseButton-secondary"]:disabled,
form [data-testid="baseButton-primary"]:disabled {{
  opacity: .6 !important;
  color: #ffffff !important;
}}
.dark-mode .stButton > button,
.dark-mode .stDownloadButton > button,
.dark-mode form [data-testid="baseButton-secondary"],
.dark-mode form [data-testid="baseButton-primary"] {{
  background: var(--tfml-btn-bg) !important;
  color: #ffffff !important;
  border-color: var(--tfml-btn-border) !important;
}}

/* File uploader: style button and hide media-capture affordances */
[data-testid="stFileUploaderDropzone"] + div button,
[data-testid="stFileUploader"] button {{
  background: var(--tfml-btn-bg) !important;
  color: #ffffff !important;
  border: 1px solid var(--tfml-btn-border) !important;
  border-radius: 10px !important;
  font-weight: 700 !important;
}}
[data-testid="stFileUploaderDropzone"] + div button:hover,
[data-testid="stFileUploader"] button:hover {{
  background: var(--tfml-btn-bg-hover) !important;
  border-color: var(--tfml-btn-bg-hover) !important;
}}

@media (max-width: 600px) {{
  .kpi {{ padding:10px; }}
  .kpi .value {{ font-size:1.4rem; }}
  .header .title {{ font-size:20px; }}
}}
</style>
""", unsafe_allow_html=True)

# ======================================
# DB LAYER
# ======================================
def init_db():
    conn = sqlite3.connect(TENDERS_DB)
    c = conn.cursor()
    c.execute("""
    CREATE TABLE IF NOT EXISTS tenders (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT,
        org TEXT,
        sector TEXT,
        deadline TEXT,
        description TEXT,
        status TEXT,
        score REAL,
        assignee TEXT,
        drafts TEXT
    )
    """)
    conn.commit()
    conn.close()

def load_rows():
    try:
        conn = sqlite3.connect(TENDERS_DB)
        c = conn.cursor()
        c.execute("SELECT * FROM tenders")
        rows = [{
            "id": r[0], "title": r[1], "org": r[2], "sector": r[3],
            "deadline": r[4], "description": r[5], "status": r[6],
            "score": r[7], "assignee": r[8],
            "drafts": json.loads(r[9]) if r[9] else []
        } for r in c.fetchall()]
        conn.close()
        return rows
    except Exception as e:
        st.error(f"Error loading tenders: {e}")
        return []

def save_row(tender):
    try:
        conn = sqlite3.connect(TENDERS_DB)
        c = conn.cursor()
        c.execute("""
        INSERT OR REPLACE INTO tenders (id, title, org, sector, deadline, description, status, score, assignee, drafts)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            tender.get("id"), tender.get("title"), tender.get("org"), tender.get("sector"),
            tender.get("deadline"), tender.get("description"), tender.get("status"),
            tender.get("score", 0.0), tender.get("assignee", ""), json.dumps(tender.get("drafts", []))
        ))
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Error saving tender: {e}")

def delete_row(tender_id):
    try:
        conn = sqlite3.connect(TENDERS_DB)
        c = conn.cursor()
        c.execute("DELETE FROM tenders WHERE id = ?", (tender_id,))
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Error deleting tender: {e}")

init_db()

# ======================================
# UTIL / DOCS
# ======================================
EOI_TMPL = """Dear {recipient},

Total Facilities Management Limited (TFML) is pleased to express interest in the opportunity titled "{title}". With a strong track record delivering {sector_desc}, our team is positioned to meet your outcomes on quality, compliance and timelines.

Scope alignment (summary):
{summary}

TFML offers certified personnel, SLAs, HSE compliance and proven delivery for public and private estates nationwide. We welcome the opportunity to submit full technical and financial proposals upon request.

Sincerely,
TFML Bid Office
"""

def _safe_date(s):
    try:
        return datetime.strptime(s, "%Y-%m-%d").date()
    except Exception:
        return None

def _suggest_email(org: str) -> str:
    org = (org or "").lower()
    if "mtn" in org: return "procurement@mtn.com"
    if "fcta" in org: return "procurement@fcta.gov.ng"
    if "faan" in org: return "procurement@faan.gov.ng"
    if "ifma" in org: return "secretariat@ifma.org.ng"
    if "aatc" in org or "afrex" in org: return "procurement@afreximbank.com"
    if "nibss" in org: return "tenders@nibss-plc.com"
    return "procurement@buyer.ng"

def write_docx_from_draft(draft: dict, filename_hint: str) -> str:
    """Generate a DOCX from a draft dict (subject/body) and save under eois/."""
    safe_fn = filename_hint[:60].replace(" ", "_")
    path = EOIS / f"{safe_fn}.docx"
    doc = Document()
    doc.add_heading(draft.get("subject", "Draft Response"), level=1)
    for meta in (f"To: {draft.get('to','')}",
                 f"CC: {draft.get('cc','')}",
                 f"Contract Value (₦): {draft.get('value','')}"):
        doc.add_paragraph(meta)
    doc.add_paragraph("")  # spacer
    body = draft.get("body") or "—"
    for line in body.split("\n"):
        doc.add_paragraph(line)
    doc.save(path)
    return str(path)

def ai_summarize(description):
    return f"Summary: {description[:180]}..."  # placeholder for real LLM

# ======================================
# SEED 6 SAMPLE TENDERS + FIRST DRAFT
# ======================================
def seed_sample_data_if_empty():
    rows_now = load_rows()
    if rows_now: return rows_now

    today = date.today()
    samples = [
        ("IFMA Abuja Secretariat FM Services", "IFMA Nigeria", "Facilities Management", 6,
         "PPP maintenance, SLA reporting & helpdesk for IFMA Secretariat, Abuja.", "Draft", "bids@tfml.ng"),
        ("AATC HQ Janitorial & Waste Management", "Afreximbank AATC", "Facilities Management", 12,
         "Janitorial, pest, waste services for AATC HQ; ISO docs & quarterly deep-clean.", "Submitted", "enoch@tfml.ng"),
        ("Wuse District Streetlighting Retrofit", "FCTA", "Energy", 3,
         "LED retrofit + solar hybridization for Wuse district roads; energy audit + M&V.", "Pending", "femi@tfml.ng"),
        ("MTN Regional Hub M&E Maintenance", "MTN Nigeria", "Construction", 20,
         "HVAC, power, fire, gen maintenance + CMMS reporting; 24/7 response.", "Draft", "greg@tfml.ng"),
        ("Airport Concourse Cleaning & Consumables", "FAAN", "Facilities Management", 9,
         "Terminal cleaning, restrooms, touchpoints; IoT counters & predictive supply.", "Submitted", "bids@tfml.ng"),
        ("Data Centre Critical Environment FM", "NIBSS", "Facilities Management", 1,
         "Tier-III: chilled water, precision cooling, UPS, suppression; 15-min response.", "Draft", "ops@tfml.ng"),
    ]

    for i, (title, org, sector, days_out, desc, status, assignee) in enumerate(samples, start=1):
        tender = {
            "id": i, "title": title, "org": org, "sector": sector,
            "deadline": (today + timedelta(days=days_out)).strftime("%Y-%m-%d"),
            "description": desc, "status": status, "score": 0.0,
            "assignee": assignee, "drafts": []
        }
        # Initial draft response
        initial = {
            "id": f"{i}:1",
            "type": "EOI",
            "version": 1,
            "status": "Draft",          # Draft -> Ready -> Sent -> Submitted
            "to": _suggest_email(org),
            "cc": "bids@tfml.ng",
            "subject": f"Expression of Interest — {title}",
            "value": "",
            "body": EOI_TMPL.format(
                recipient="Procurement Team", title=title,
                sector_desc=sector.lower(), summary=desc
            ),
            "attachments": [],
            "file": "",
            "last_updated": datetime.now().isoformat(timespec="seconds")
        }
        tender["drafts"] = [initial]
        save_row(tender)

    return load_rows()

# ======================================
# FAKE EMAIL SENDER (placeholder)
# ======================================
def send_email(recipient, subject, body, attachment_paths=None, cc=None):
    cc_text = f" (cc: {cc})" if cc else ""
    st.success(f"Email queued to **{recipient}**{cc_text} with subject **{subject}**.")

# ======================================
# HEADER
# ======================================
def logo_header():
    cols = st.columns([0.12, 0.88])
    with cols[0]:
        if LOGO_PATH.exists():
            st.image(Image.open(LOGO_PATH), use_container_width=True)
        else:
            st.write("**TFML**")
    with cols[1]:
        st.markdown(f"<div class='header'><div class='title'>Agentic AI Console</div></div>", unsafe_allow_html=True)
        st.markdown(f"<span style='color:{ACCENT};opacity:.9;'>It’s all about you… • one-click drafting • faster BD • higher win rate</span>", unsafe_allow_html=True)

logo_header()

# ======================================
# LOAD DATA + SOON DUE NOTICES
# ======================================
rows = seed_sample_data_if_empty()

def render_deadline_notices(rows, days=3):
    today = datetime.today().date()
    soon = today + timedelta(days=days)
    for r in rows:
        d = _safe_date(r.get("deadline"))
        if d and d <= soon:
            st.warning(f"⚠️ Tender '{r.get('title','Untitled')}' is due on {d.strftime('%Y-%m-%d')}!")
render_deadline_notices(rows, days=3)

# ======================================
# DASHBOARD HELPERS
# ======================================
def compute_dashboard_metrics(rows):
    today = datetime.today().date()
    in_7 = today + timedelta(days=7)
    in_3 = today + timedelta(days=3)

    total = len(rows)
    deadlines = [(_safe_date(r.get("deadline")), r) for r in rows]
    overdue = [r for d, r in deadlines if d and d < today and r.get("status") not in ("Awarded", "Won", "Lost")]
    due3 = [r for d, r in deadlines if d and today <= d <= in_3]
    due7 = [r for d, r in deadlines if d and today <= d <= in_7]
    drafts = [r for r in rows if r.get("status") == "Draft"]
    inflight = [r for r in rows if r.get("status") in ("Submitted", "Pending")]
    awarded = [r for r in rows if r.get("status") in ("Awarded", "Won")]
    decided = [r for r in rows if r.get("status") in ("Awarded", "Won", "Lost")]
    win_rate = round((len(awarded) / len(decided) * 100.0), 1) if decided else 0.0

    by_assignee = {}
    for r in rows:
        a = (r.get("assignee") or "Unassigned").strip() or "Unassigned"
        by_assignee[a] = by_assignee.get(a, 0) + 1

    next30 = []
    for d, r in deadlines:
        if d and today <= d <= (today + timedelta(days=30)):
            next30.append(d)
    df_next30 = (pd.Series(next30, name="date").value_counts().rename_axis("date").reset_index(name="tenders").sort_values("date")
                 if next30 else pd.DataFrame(columns=["date", "tenders"]))

    feed = []
    for r in rows:
        for d in r.get("drafts", []):
            feed.append({
                "when": os.path.getmtime(d["file"]) if d.get("file") and os.path.exists(d["file"]) else time.time(),
                "tender": r.get("title", "Untitled"),
                "type": d.get("type", "Doc"),
                "file": os.path.basename(d.get("file") or ""),
                "version": d.get("version", 1),
                "status": d.get("status", "")
            })
    feed_df = pd.DataFrame(feed)
    if not feed_df.empty:
        feed_df["when"] = pd.to_datetime(feed_df["when"], unit="s")
        feed_df = feed_df.sort_values("when", ascending=False)

    return {
        "total": total, "overdue": len(overdue), "due3": len(due3), "due7": len(due7),
        "drafts": len(drafts), "inflight": len(inflight), "awarded": len(awarded), "win_rate": win_rate,
        "assignee_counts": by_assignee, "deadline_30": df_next30, "activity": feed_df
    }

# ======================================
# DRAFT HELPERS
# ======================================
def new_draft_response_for_tender(tender: dict, kind="EOI"):
    """Create and attach a new draft response with sensible defaults."""
    next_version = (max([d.get("version", 0) for d in tender.get("drafts", [])]) + 1) if tender.get("drafts") else 1
    draft_id = f"{tender['id']}:{next_version}"
    draft = {
        "id": draft_id,
        "type": kind,
        "version": next_version,
        "status": "Draft",          # Draft -> Ready -> Sent -> Submitted
        "to": _suggest_email(tender.get("org")),
        "cc": "bids@tfml.ng",
        "subject": f"{'Proposal' if kind!='EOI' else 'Expression of Interest'} — {tender.get('title','')}",
        "value": "",                # ₦
        "body": EOI_TMPL.format(
            recipient="Procurement Team",
            title=tender.get("title","Untitled"),
            sector_desc=tender.get("sector","Facilities Management").lower(),
            summary=tender.get("description","—"),
        ),
        "attachments": [],
        "file": "",
        "last_updated": datetime.now().isoformat(timespec="seconds")
    }
    tender["drafts"] = (tender.get("drafts") or []) + [draft]
    save_row(tender)
    return draft

def validate_email_list(s: str) -> bool:
    if not s: return True
    emails = [e.strip() for e in s.split(",") if e.strip()]
    simple = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
    return all(simple.match(e) for e in emails)

# ======================================
# TABS
# ======================================
tab_dash, tab_tenders, tab_drafts, tab_settings = st.tabs(["Dashboard", "Tenders", "Drafts", "Settings"])

# ======================================
# DASHBOARD
# ======================================
with tab_dash:
    st.markdown("#### Executive Overview")
    m = compute_dashboard_metrics(rows)

    c1, c2, c3, c4, c5, c6 = st.columns(6)
    with c1: st.markdown(f"<div class='kpi'><div class='label'>Total</div><div class='value'>{m['total']}</div><div class='sub'>All notices</div></div>", unsafe_allow_html=True)
    with c2: st.markdown(f"<div class='kpi'><div class='label'>Overdue</div><div class='value'>{m['overdue']}</div><div class='sub'>Past deadline</div></div>", unsafe_allow_html=True)
    with c3: st.markdown(f"<div class='kpi'><div class='label'>Due in 3 days</div><div class='value'>{m['due3']}</div><div class='sub'>Immediate action</div></div>", unsafe_allow_html=True)
    with c4: st.markdown(f"<div class='kpi'><div class='label'>Due in 7 days</div><div class='value'>{m['due7']}</div><div class='sub'>Upcoming</div></div>", unsafe_allow_html=True)
    with c5: st.markdown(f"<div class='kpi'><div class='label'>In Flight</div><div class='value'>{m['inflight']}</div><div class='sub'>Submitted/Pending</div></div>", unsafe_allow_html=True)
    with c6: st.markdown(f"<div class='kpi'><div class='label'>Win rate</div><div class='value'>{m['win_rate']}%</div><div class='sub'>Awards</div></div>", unsafe_allow_html=True)

    st.markdown("---")
    left, right = st.columns([0.6, 0.4])

    with left:
        if rows:
            df = pd.DataFrame(rows)
            df["deadline_dt"] = pd.to_datetime(df["deadline"], errors="coerce")

            st.markdown("##### Tenders by Sector")
            sector_chart = alt.Chart(df).mark_bar(cornerRadiusTopLeft=4, cornerRadiusTopRight=4).encode(
                x=alt.X('sector:N', sort='-y', title=''),
                y=alt.Y('count():Q', title='Tenders'),
                tooltip=['sector', 'count()'],
                color=alt.Color('sector:N', scale=alt.Scale(scheme='category10'), legend=None)
            ).properties(height=220, background='transparent')
            st.altair_chart(sector_chart, use_container_width=True)

            st.markdown("##### Pipeline Status")
            donut = alt.Chart(df).mark_arc(innerRadius=70).encode(
                theta=alt.Theta("count():Q"),
                color=alt.Color("status:N", scale=alt.Scale(scheme='category10')),
                tooltip=["status", "count()"]
            ).properties(height=240)
            st.altair_chart(donut, use_container_width=True)

            st.markdown("##### Deadline Load (Next 30 Days)")
            dl = m["deadline_30"]
            if not dl.empty:
                area = alt.Chart(dl).mark_area(opacity=0.6).encode(
                    x=alt.X("date:T", title="Date"),
                    y=alt.Y("tenders:Q", title="Count"),
                    tooltip=["date:T", "tenders:Q"]
                ).properties(height=220)
                st.altair_chart(area, use_container_width=True)
            else:
                st.info("No deadlines in the next 30 days.")
        else:
            st.info("No tenders yet. Add a few to unlock insights.")

    with right:
        st.markdown("##### Smart Alerts")
        if m["overdue"] > 0:
            st.error(f"{m['overdue']} tender(s) are overdue. Prioritise these now.")
        elif m["due3"] > 0:
            st.warning(f"{m['due3']} tender(s) due in 3 days.")
        elif m["due7"] > 0:
            st.warning(f"{m['due7']} tender(s) due in 7 days.")
        else:
            st.success("All clear. No urgent deadlines.")

        st.markdown("##### Workload by Assignee")
        ass_df = pd.DataFrame([{"Assignee": k, "Tenders": v} for k, v in m["assignee_counts"].items()])
        if not ass_df.empty:
            bar = alt.Chart(ass_df).mark_bar(cornerRadiusTopLeft=4, cornerRadiusTopRight=4).encode(
                x=alt.X("Assignee:N", sort='-y', title=''),
                y=alt.Y("Tenders:Q", title="Count"),
                tooltip=["Assignee", "Tenders"]
            ).properties(height=220)
            st.altair_chart(bar, use_container_width=True)
        else:
            st.info("No assignments yet.")

        st.markdown("##### Top Upcoming Deadlines")
        soon = []
        for r in rows:
            d = _safe_date(r.get("deadline"))
            if d:
                soon.append({"Deadline": d.strftime("%Y-%m-%d"), "Title": r.get("title",""),
                             "Status": r.get("status",""), "Assignee": r.get("assignee","")})
        soon = sorted(soon, key=lambda x: x["Deadline"])[:10]
        if soon:
            st.dataframe(pd.DataFrame(soon), use_container_width=True, hide_index=True)
        else:
            st.info("No upcoming deadlines found.")

    st.markdown("---")
    st.markdown("#### Activity Feed")
    act = m["activity"]
    if not act.empty:
        af = act[["when","tender","type","version","file","status"]].rename(columns={
            "when":"Time","tender":"Tender","type":"Doc","version":"v","file":"File","status":"Status"
        })
        st.dataframe(af.head(15), use_container_width=True, hide_index=True)
    else:
        st.caption("No document activity yet. Generate a draft response to see activity here.")

# ======================================
# TENDERS
# ======================================
with tab_tenders:
    st.markdown("### Manage Tenders")

    # Filters
    colf1, colf2, colf3, colf4 = st.columns([0.35, 0.2, 0.25, 0.2])
    with colf1:
        search = st.text_input("Search title or buyer", placeholder="e.g., 'airport' or 'FAAN'")
    with colf2:
        all_statuses = ["Draft", "Submitted", "Pending", "Awarded", "Won", "Lost"]
        status_filter = st.multiselect("Status", all_statuses, default=["Draft", "Submitted", "Pending"])
    with colf3:
        sectors = sorted({r.get("sector","") for r in rows if r.get("sector")}) or ["Facilities Management","Construction","Energy","Other"]
        sector_filter = st.multiselect("Sector", sectors, default=sectors)
    with colf4:
        today = datetime.today().date()
        start_date = st.date_input("From", today - timedelta(days=14))
        end_date = st.date_input("To", today + timedelta(days=60))

    # Natural language
    nl_query = st.text_input("Ask about tenders (e.g., 'show overdue', 'due this week')")
    def process_nl(q, rs):
        if not q: return rs
        q = q.lower().strip()
        if "overdue" in q: return [r for r in rs if _safe_date(r.get("deadline")) and _safe_date(r["deadline"]) < today]
        if "due this week" in q: return [r for r in rs if _safe_date(r.get("deadline")) and _safe_date(r["deadline"]) <= (today + timedelta(days=7))]
        return rs
    filtered = process_nl(nl_query, rows)

    def _match(r):
        t = (r.get("title","") + " " + r.get("org","")).lower()
        d = _safe_date(r.get("deadline"))
        in_range = (d is None) or (start_date <= d <= end_date)
        return (search.lower() in t) and (r.get("status") in status_filter) and (r.get("sector") in sector_filter) and in_range
    filtered = [r for r in filtered if _match(r)]

    sub_list, sub_kanban, sub_calendar = st.tabs(["List", "Kanban", "Calendar"])

    # -------- List View --------
    with sub_list:
        if filtered:
            by_assignee = {}
            for r in filtered:
                a = (r.get("assignee") or "Unassigned").strip() or "Unassigned"
                by_assignee[a] = by_assignee.get(a, 0) + 1
            if by_assignee:
                chips = " ".join([f"<span class='pill'>{a}: {n}</span>" for a,n in by_assignee.items()])
                st.markdown(chips, unsafe_allow_html=True)
            st.write("")

            for r in filtered:
                row_cols = st.columns([0.6, 0.2, 0.2])
                with row_cols[0]:
                    st.markdown(f"**{r['title']}**  \n_{r['org']}_")
                with row_cols[1]:
                    st.markdown(f"**Deadline:** {r['deadline']}  \n**Status:** {r['status']}")
                with row_cols[2]:
                    st.markdown(f"**Sector:** {r['sector']}  \n**Assignee:** {r.get('assignee','')}")

                with st.expander("Details / Actions"):
                    st.write(f"**AI Summary:** {ai_summarize(r.get('description',''))}")

                    c1, c2, c3, c4 = st.columns(4)
                    with c1:
                        if st.button("Generate Draft Response", key=f"gen_{r['id']}"):
                            d = new_draft_response_for_tender(r, kind="EOI")
                            st.success(f"Draft created (v{d['version']}). Edit it in **Drafts** tab.")
                    with c2:
                        new_status = st.selectbox("Update Status", all_statuses, index=all_statuses.index(r.get("status","Draft")), key=f"ust_{r['id']}")
                        if new_status != r.get("status"):
                            r["status"] = new_status
                            save_row(r)
                            st.info("Status updated.")
                    with c3:
                        new_score = st.slider("Fit Score", 0.0, 100.0, float(r.get("score") or 0.0), 1.0, key=f"scr_{r['id']}")
                        if new_score != r.get("score"):
                            r["score"] = float(new_score)
                            save_row(r)
                    with c4:
                        if st.button("Delete", key=f"del_{r['id']}"):
                            delete_row(r["id"])
                            st.success("Tender deleted.")
                            try: st.rerun()
                            except Exception: st.experimental_rerun()
        else:
            st.info("No tenders match the filters.")

    # -------- Kanban --------
    with sub_kanban:
        cols = st.columns(5)
        lanes = [("Draft", cols[0]), ("Submitted", cols[1]), ("Pending", cols[2]), ("Won", cols[3]), ("Lost", cols[4])]
        for status, col in lanes:
            with col:
                st.markdown(f"**{status}**")
                lane_items = [r for r in filtered if r.get("status")==status]
                if not lane_items: st.caption("—")
                for r in lane_items:
                    st.markdown(
                        f"<div class='card'><strong>{r['title']}</strong><br>"
                        f"<span class='pill'>{r['sector']}</span> <span class='pill'>Due: {r['deadline']}</span><br>"
                        f"<small>{r.get('org','')}</small></div>", unsafe_allow_html=True
                    )
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("Draft Response", key=f"kgen_{r['id']}"):
                            d = new_draft_response_for_tender(r, kind="EOI")
                            st.success("Draft created. Edit it in Drafts.")
                    with c2:
                        nxt_opts = [s for s in ["Draft","Submitted","Pending","Won","Lost"] if s != status]
                        nxt = st.selectbox("Move to", nxt_opts, key=f"kan_mv_{r['id']}")
                        if st.button("Move", key=f"kan_btn_{r['id']}"):
                            r["status"] = nxt
                            save_row(r)
                            st.info(f"Moved to {nxt}")
                            try: st.rerun()
                            except Exception: st.experimental_rerun()

    # -------- Calendar --------
    with sub_calendar:
        if filtered:
            dfc = pd.DataFrame(filtered)
            dfc["deadline_dt"] = pd.to_datetime(dfc["deadline"], errors="coerce")
            cal = alt.Chart(dfc.dropna(subset=["deadline_dt"])).mark_circle(size=110).encode(
                x=alt.X("deadline_dt:T", title="Deadline"),
                y=alt.Y("sector:N", title="Sector"),
                color=alt.Color("status:N", scale=alt.Scale(scheme="category10")),
                tooltip=["title","org","deadline","status","assignee"]
            ).properties(height=320)
            st.altair_chart(cal, use_container_width=True)
        else:
            st.info("Nothing to plot.")

# ======================================
# DRAFTS WORKSPACE
# ======================================
with tab_drafts:
    st.markdown("### Drafts Workspace")

    # Flatten drafts
    draft_rows = []
    for r in rows:
        for i, d in enumerate(r.get("drafts", [])):
            draft_rows.append({
                "DraftID": d.get("id") or f"{r['id']}:{i+1}",
                "TenderID": r["id"],
                "Tender": r.get("title",""),
                "Buyer": r.get("org",""),
                "Type": d.get("type","EOI"),
                "Version": d.get("version",1),
                "Status": d.get("status","Draft"),
                "Value(₦)": d.get("value",""),
                "To": d.get("to",""),
                "CC": d.get("cc",""),
                "Subject": d.get("subject",""),
                "Last Updated": d.get("last_updated",""),
                "_body": d.get("body",""),
                "_file": d.get("file",""),
                "_attachments": d.get("attachments",[])
            })
    df_drafts = pd.DataFrame(draft_rows)

    # Metrics
    if not df_drafts.empty:
        def _to_float(v):
            try: return float(str(v).replace(",",""))
            except: return 0.0
        total_value = df_drafts["Value(₦)"].apply(_to_float).sum()
        colm1, colm2, colm3 = st.columns(3)
        with colm1:
            st.markdown(f"<div class='kpi'><div class='label'>Total Drafts</div><div class='value'>{len(df_drafts)}</div><div class='sub'>All types</div></div>", unsafe_allow_html=True)
        with colm2:
            ready = (df_drafts["Status"]=="Ready").sum()
            sent = (df_drafts["Status"]=="Sent").sum()
            submitted = (df_drafts["Status"]=="Submitted").sum()
            st.markdown(f"<div class='kpi'><div class='label'>Status</div><div class='value'>{ready} Ready</div><div class='sub'>{sent} Sent • {submitted} Submitted</div></div>", unsafe_allow_html=True)
        with colm3:
            st.markdown(f"<div class='kpi'><div class='label'>Total Value</div><div class='value'>₦{total_value:,.0f}</div><div class='sub'>Across drafts</div></div>", unsafe_allow_html=True)

    # Filters + selection
    colf1, colf2, colf3 = st.columns([0.45, 0.25, 0.3])
    with colf1:
        q = st.text_input("Search (Tender / Buyer / Subject)")
    with colf2:
        status_opts = ["Draft","Ready","Sent","Submitted"]
        f_status = st.multiselect("Filter Status", status_opts, default=status_opts)
    with colf3:
        type_opts = ["EOI","Proposal"]
        f_type = st.multiselect("Type", type_opts, default=type_opts)

    def _match_draft(row):
        t = f"{row['Tender']} {row['Buyer']} {row['Subject']}".lower()
        if q and q.lower() not in t: return False
        if row["Status"] not in f_status: return False
        if row["Type"] not in f_type: return False
        return True

    if not df_drafts.empty:
        view_df = df_drafts[df_drafts.apply(_match_draft, axis=1)].copy()
    else:
        view_df = pd.DataFrame(columns=["DraftID","Tender","Buyer","Type","Version","Status","Value(₦)","To","CC","Subject","Last Updated"])

    col_list, col_edit = st.columns([0.55, 0.45])

    # -------- Left: list of drafts
    with col_list:
        if view_df.empty:
            st.info("No drafts yet. Generate one from the Tenders page.")
        else:
            show_cols = ["DraftID","Tender","Buyer","Type","Version","Status","Value(₦)","To","Subject","Last Updated"]
            st.dataframe(view_df[show_cols], use_container_width=True, hide_index=True)
            draft_ids = view_df["DraftID"].tolist()
            selected_id = st.selectbox("Select a draft to edit", draft_ids)

    # -------- Right: editor
    with col_edit:
        if not view_df.empty:
            row = view_df[view_df["DraftID"]==selected_id].iloc[0]
            t = next((r for r in rows if r["id"]==row["TenderID"]), None)
            d_index = None
            d_obj = None
            if t:
                for idx, d in enumerate(t.get("drafts", [])):
                    if (d.get("id")==row["DraftID"]) or (d.get("version")==row["Version"]):
                        d_index = idx; d_obj = d; break

            st.markdown(f"#### Edit Draft — {row['Tender']} (v{row['Version']})")
            with st.form("edit_draft_form"):
                col_a, col_b = st.columns(2)
                with col_a:
                    to = st.text_input("To (emails, comma-separated)", value=row["To"])
                    cc = st.text_input("CC (comma-separated)", value=row["CC"])
                    value_str = st.text_input("Contract Value (₦)", value=str(row["Value(₦)"]))
                    status = st.selectbox("Status", ["Draft","Ready","Sent","Submitted"],
                                          index=["Draft","Ready","Sent","Submitted"].index(row["Status"]))
                with col_b:
                    subject = st.text_input("Subject", value=row["Subject"])
                    draft_type = st.selectbox("Type", ["EOI","Proposal"], index=["EOI","Proposal"].index(row["Type"]))
                    # Restrict to common doc/image types to avoid audio control
                    attach = st.file_uploader(
                        "Add attachment(s)",
                        accept_multiple_files=True,
                        type=["pdf","doc","docx","ppt","pptx","xls","xlsx","csv","txt","rtf","zip","png","jpg","jpeg"]
                    )

                body = st.text_area("Body", value=row["_body"], height=280)

                save = st.form_submit_button("💾 Save Changes")
                if save:
                    if not validate_email_list(to) or not validate_email_list(cc):
                        st.error("Please enter valid email addresses (comma-separated).")
                    else:
                        if t and d_obj is not None:
                            d_obj.update({
                                "to": to, "cc": cc, "value": value_str, "subject": subject,
                                "type": draft_type, "body": body, "status": status,
                                "last_updated": datetime.now().isoformat(timespec="seconds")
                            })
                            # Save attachments
                            if attach:
                                saved = []
                                for f in attach:
                                    savep = EOIS / f.name
                                    with open(savep, "wb") as out:
                                        out.write(f.read())
                                    saved.append(str(savep))
                                d_obj["attachments"] = list(set((d_obj.get("attachments") or []) + saved))
                            save_row(t)
                            st.success("Draft updated.")

            # actions row
            ca, cb, cc2, cd, ce = st.columns(5)
            with ca:
                if st.button("📥 Download DOCX"):
                    filename_hint = f"{row['Tender']}_v{row['Version']}"
                    file_path = write_docx_from_draft(d_obj, filename_hint)
                    d_obj["file"] = file_path
                    d_obj["last_updated"] = datetime.now().isoformat(timespec="seconds")
                    save_row(t)
                    with open(file_path, "rb") as f:
                        st.download_button("Download file", f, file_name=Path(file_path).name, use_container_width=True)
            with cb:
                if st.button("🧬 Duplicate (Version +1)"):
                    new_ver = (max([d.get("version",0) for d in t.get("drafts",[])])+1) if t.get("drafts") else 1
                    new_id = f"{t['id']}:{new_ver}"
                    clone = dict(d_obj)
                    clone["version"] = new_ver
                    clone["id"] = new_id
                    clone["status"] = "Draft"
                    clone["file"] = ""
                    clone["last_updated"] = datetime.now().isoformat(timespec="seconds")
                    t["drafts"].append(clone)
                    save_row(t)
                    st.success(f"Duplicated as v{new_ver}.")
                    try: st.rerun()
                    except Exception: st.experimental_rerun()
            with cc2:
                if st.button("✅ Mark as Submitted"):
                    d_obj["status"] = "Submitted"
                    d_obj["last_updated"] = datetime.now().isoformat(timespec="seconds")
                    save_row(t)
                    st.success("Marked as Submitted.")
            with cd:
                if st.button("✉️ Send Email"):
                    if not d_obj.get("to"):
                        st.error("Enter a recipient email first.")
                    else:
                        send_email(d_obj.get("to"), d_obj.get("subject"), d_obj.get("body"),
                                   d_obj.get("attachments"), cc=d_obj.get("cc"))
                        d_obj["status"] = "Sent"
                        d_obj["last_updated"] = datetime.now().isoformat(timespec="seconds")
                        save_row(t)
            with ce:
                if st.button("🗑️ Delete Draft"):
                    if t and d_index is not None:
                        t["drafts"].pop(d_index)
                        save_row(t)
                        st.warning("Draft deleted.")
                        try: st.rerun()
                        except Exception: st.experimental_rerun()

# ======================================
# SETTINGS
# ======================================
with tab_settings:
    st.markdown("### Settings")
    st.session_state["default_recipient"] = st.text_input("Default Recipient", value=st.session_state.get("default_recipient","Procurement Team"))
    st.session_state["bid_email"] = st.text_input("Bid Office Email", value=st.session_state.get("bid_email","bids@tfml.ng"))
    st.session_state["bid_phone"] = st.text_input("Bid Office Phone", value=st.session_state.get("bid_phone","+234-XXX-XXXX"))
    theme = st.selectbox("Theme", ["Light","Dark"])
    if theme == "Dark":
        st.markdown("<script>document.body.classList.add('dark-mode');</script>", unsafe_allow_html=True)
    else:
        st.markdown("<script>document.body.classList.remove('dark-mode');</script>", unsafe_allow_html=True)
    st.caption("Changes save automatically.")
